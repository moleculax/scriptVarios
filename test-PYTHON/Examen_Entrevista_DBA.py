from docx import Document

# Crear documento
doc = Document()
doc.add_heading('Examen Simulado de Entrevista para Administrador DBA', 0)

# Secciones con preguntas y respuestas
contenido = {
    "Sección 1: Conceptos Generales": [
        ("Explica la diferencia entre OLTP y OLAP.",
         "OLTP (Online Transaction Processing) está orientado a transacciones rápidas y consistentes, con muchas operaciones pequeñas (ej. sistemas bancarios). "
         "OLAP (Online Analytical Processing) está orientado al análisis de grandes volúmenes de datos, con consultas complejas y agregaciones (ej. sistemas de BI)."),
        
        ("¿Qué es un índice en una base de datos? ¿Qué ventajas y desventajas tiene?",
         "Un índice es una estructura de datos que acelera la búsqueda de registros. Ventajas: mejora la velocidad de consultas. "
         "Desventajas: consume espacio en disco y puede ralentizar inserciones/actualizaciones."),
        
        ("¿Qué diferencia existe entre clúster de alta disponibilidad y replicación?",
         "El clúster de alta disponibilidad asegura continuidad de servicio con nodos redundantes, mientras que la replicación copia datos entre servidores para balanceo de carga o recuperación."),
        
        ("Explica la diferencia entre commit y rollback.",
         "Commit confirma de forma permanente los cambios en la base de datos, mientras que rollback revierte los cambios no confirmados."),
        
        ("¿Qué es un plan de ejecución y cómo lo analizarías?",
         "Es la representación de cómo el motor ejecutará una consulta. Se analiza buscando operaciones costosas (scans completos, joins ineficientes) y se proponen índices o reescritura de consultas.")
    ],

    "Sección 2: SQL Práctico": [
        ("Escribe una consulta SQL que devuelva el segundo salario más alto de una tabla empleados(id, nombre, salario).",
         "SELECT MAX(salario) FROM empleados WHERE salario < (SELECT MAX(salario) FROM empleados);"),
        
        ("Dada la tabla ventas(id, fecha, monto), escribe una consulta que muestre el monto total vendido por mes en 2024.",
         "SELECT MONTH(fecha) AS mes, SUM(monto) AS total FROM ventas WHERE YEAR(fecha)=2024 GROUP BY MONTH(fecha);"),
        
        ("Optimiza la siguiente consulta: SELECT * FROM pedidos WHERE YEAR(fecha) = 2024;",
         "Es mejor usar rangos para aprovechar índices: SELECT * FROM pedidos WHERE fecha BETWEEN '2024-01-01' AND '2024-12-31';"),
        
        ("Dada la tabla clientes(id, nombre, email), escribe una consulta que encuentre los emails duplicados.",
         "SELECT email, COUNT(*) FROM clientes GROUP BY email HAVING COUNT(*) > 1;")
    ],

    "Sección 3: Administración y Mantenimiento": [
        ("¿Cómo planificarías una estrategia de backup para una base de datos de 1 TB que debe estar disponible 24/7?",
         "Implementaría backups completos semanales, diferenciales diarios y logs de transacciones cada pocos minutos, además de un servidor en standby para recuperación rápida."),
        
        ("¿Qué diferencia hay entre full backup, differential backup y transaction log backup?",
         "Full backup: copia completa. Differential: copia de los cambios desde el último full. Transaction log: registra todas las transacciones desde el último backup, permitiendo recuperación punto en el tiempo."),
        
        ("¿Cómo monitorearías el uso de espacio en disco de una base de datos en producción?",
         "Mediante alertas del sistema operativo, vistas de administración del motor (ej. sys.dm_db_partition_stats en SQL Server), y herramientas de monitoreo."),
        
        ("¿Qué pasos seguirías si un servidor presenta alto consumo de CPU debido a consultas pesadas?",
         "Identificar consultas con mayor consumo (mediante profiler o slow query log), revisar planes de ejecución, optimizar índices y considerar particionamiento.")
    ],

    "Sección 4: Performance Tuning": [
        ("¿Qué parámetros del motor sueles revisar para mejorar rendimiento?",
         "Tamaño de buffer pool, configuración de cache, paralelismo de consultas, estadísticas actualizadas, índices adecuados."),
        
        ("¿Qué es el Index Seek vs Index Scan en un plan de ejecución?",
         "Index Seek accede directamente a las filas necesarias (eficiente). Index Scan recorre todo el índice (menos eficiente)."),
        
        ("Explica cómo detectar queries lentas en MySQL, Oracle y SQL Server.",
         "MySQL: slow_query_log. Oracle: AWR reports. SQL Server: DMV sys.dm_exec_query_stats o Profiler.")
    ],

    "Sección 5: Seguridad": [
        ("¿Cómo implementarías el principio de mínimo privilegio en un sistema con múltiples aplicaciones?",
         "Asignando roles específicos con permisos mínimos necesarios, evitando que aplicaciones usen usuarios con permisos de administrador."),
        
        ("¿Qué medidas aplicarías para proteger los datos en tránsito y en reposo?",
         "Cifrado TLS para datos en tránsito y cifrado de disco/base de datos (TDE, AES) para datos en reposo."),
        
        ("¿Qué es el SQL Injection y cómo lo previenes desde la administración de la base de datos?",
         "Es una técnica que manipula consultas mediante entradas maliciosas. Se previene usando consultas parametrizadas, validación de entradas y limitando privilegios.")
    ],

    "Sección 6: Escenarios de Resolución de Problemas": [
        ("Un usuario reporta que la base de datos está lenta. ¿Qué pasos sigues para diagnosticar el problema?",
         "1. Verificar recursos (CPU, memoria, disco). 2. Revisar queries activas. 3. Analizar índices y planes de ejecución. 4. Revisar bloqueos o deadlocks."),
        
        ("La replicación entre dos servidores ha fallado. ¿Cómo identificarías la causa y qué acciones tomarías?",
         "Revisar logs de replicación, comparar versiones de datos, reiniciar el proceso de replicación o reconfigurarlo si hay divergencias."),
        
        ("Una tabla crítica ha sido borrada accidentalmente. ¿Cómo restaurarías los datos sin perder cambios recientes?",
         "Restaurar backup completo más transaction logs hasta el punto anterior al borrado (point-in-time recovery)."),
        
        ("El espacio en disco se llenó en un servidor en producción. ¿Qué medidas urgentes tomarías?",
         "Eliminar logs viejos, mover archivos de backup a otro storage, comprimir índices/tablas y ampliar almacenamiento si es posible.")
    ]
}

# Agregar contenido al documento
for seccion, preguntas in contenido.items():
    doc.add_heading(seccion, level=1)
    for pregunta, respuesta in preguntas:
        doc.add_paragraph(f"❓ {pregunta}", style='List Bullet')
        doc.add_paragraph(f"✅ Respuesta: {respuesta}\n")

# Guardar documento
file_path = "Examen_Entrevista_DBA.docx"
doc.save(file_path)
file_path
