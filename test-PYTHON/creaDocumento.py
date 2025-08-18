from docx import Document

# Crear documento Word nuevamente (el estado se reinició)
docx_path_guide = "/mnt/data/guia_estudio_java_backend.docx"
document = Document()

document.add_heading("Guía de Estudio – Examen Backend Java", 0)

# Secciones de la guía
sections = {
    "1. Fundamentos de Java": [
        "POO: clases, objetos, herencia, polimorfismo, encapsulación.",
        "Colecciones: List, Set, Map (diferencias, cuándo usar cada una).",
        "Generics: uso de <T>.",
        "Manejo de excepciones: try-catch-finally, excepciones checked y unchecked.",
        "Java 8+: Lambdas, Streams, Optional.",
        "Multithreading y concurrencia: Thread, Runnable, ExecutorService, synchronized.",
        "Garbage Collector y memoria: stack vs heap, referencias fuertes, débiles y suaves.",
        "Ejemplo: implementar un método que filtre una lista con Streams."
    ],
    "2. Spring Framework / Spring Boot": [
        "Arquitectura de Spring: IoC, AOP, Beans.",
        "Spring Boot: autoconfiguración, starters, application.properties.",
        "Inyección de dependencias: @Autowired, @Component, @Service, @Repository.",
        "REST API: @RestController, @GetMapping, @PostMapping, ResponseEntity, validaciones con @Valid.",
        "Spring Data JPA: JpaRepository, @Query.",
        "Spring Security: autenticación, roles, JWT.",
        "Testing: @SpringBootTest, MockMvc.",
        "Ejemplo: crear un endpoint /users con operaciones CRUD."
    ],
    "3. Hibernate / JPA": [
        "Entidades (@Entity, @Table).",
        "Relaciones: @OneToMany, @ManyToMany.",
        "Ciclo de vida de una entidad: persist, merge, detach, remove.",
        "JPQL y Criteria API.",
        "Caché de Hibernate.",
        "Ejemplo: modelar entidades Usuario y Rol con relación ManyToMany."
    ],
    "4. Bases de Datos y SQL": [
        "Consultas básicas: SELECT, INSERT, UPDATE, DELETE.",
        "Joins: INNER, LEFT, RIGHT, FULL.",
        "Funciones de agregación: COUNT, SUM, AVG.",
        "Índices y claves primarias/foráneas.",
        "Transacciones (ACID).",
        "Normalización (1FN, 2FN, 3FN).",
        "Ejemplo: obtener los 5 usuarios más recientes con SQL."
    ],
    "5. Arquitectura y Buenas Prácticas": [
        "Monolito vs Microservicios.",
        "Principios SOLID.",
        "Patrones de diseño: Singleton, Factory, DAO, Repository.",
        "RESTful API buenas prácticas: status codes, versionado.",
        "Mensajería: Kafka, RabbitMQ (básico).",
        "Caching: Redis.",
        "Ejemplo: diseñar arquitectura de un sistema de pedidos con microservicios."
    ],
    "6. Herramientas y DevOps": [
        "Maven/Gradle: dependencias, perfiles.",
        "Git: ramas, merge, pull request.",
        "Docker: levantar un contenedor con BD.",
        "CI/CD: Jenkins, GitHub Actions (básico).",
        "Ejemplo: compilar y ejecutar un proyecto con Maven (mvn clean install)."
    ],
    "7. Ejercicios Prácticos Recomendados": [
        "Implementar un CRUD de usuarios con Spring Boot + JPA.",
        "Crear un servicio REST que exponga información en formato JSON.",
        "Conectar Spring Boot con una base de datos (MySQL/PostgreSQL).",
        "Implementar un login con Spring Security y JWT.",
        "Resolver ejercicios de algoritmos en Java (palíndromos, ordenar listas, buscar duplicados)."
    ],
    "Tips para el examen": [
        "Explicar conceptos con ejemplos de código.",
        "Repasar logs de errores comunes en Spring/Hibernate.",
        "Hacer un repaso práctico con un mini proyecto real.",
        "Usar LeetCode / HackerRank para algoritmos básicos en Java."
    ]
}

# Agregar secciones al documento
for title, points in sections.items():
    document.add_heading(title, level=1)
    for point in points:
        document.add_paragraph(point, style="List Bullet")

# Guardar documento
document.save(docx_path_guide)

docx_path_guide
